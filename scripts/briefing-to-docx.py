"""
briefing-to-docx.py -- Convert a daily briefing markdown file into a styled Word .docx.

Usage: python scripts/briefing-to-docx.py <input.md> <output.docx>

Handles both structured markdown (## SECTION N: headers, markdown tables) and
flat text formats. Matches the "Daily Executive Briefing" visual theme:
navy headers, blue article titles, gray sources, Word tables with shading.
"""

import sys
import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# -- Colors (match the approved template) ------------------------------------
NAVY = RGBColor(0x1F, 0x39, 0x64)
BLUE = RGBColor(0x1F, 0x5C, 0x99)
GRAY = RGBColor(0x80, 0x80, 0x80)
BLACK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xE0, 0x60, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# -- Sizes -------------------------------------------------------------------
TITLE_SZ = Pt(22)
SUBTITLE_SZ = Pt(10)
SECTION_SZ = Pt(16)
ARTICLE_SZ = Pt(12)
SOURCE_SZ = Pt(9)
BODY_SZ = Pt(10.5)
SMALL_SZ = Pt(9.5)

TABLE_HEADER_BG = "1F3964"
TABLE_ALT_BG = "F2F6FA"


def _run(p, text, size=BODY_SZ, bold=False, italic=False, color=BLACK, font='Calibri'):
    r = p.add_run(text)
    r.font.size = size
    r.font.name = font
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def _para(doc, text, size=BODY_SZ, bold=False, italic=False, color=BLACK,
          align=None, before=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = before
    if after is not None:
        p.paragraph_format.space_after = after
    _run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def _shade_cell(cell, hex_color):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): hex_color,
    })
    tcPr.append(shading)


def _add_table(doc, rows):
    """Add a styled table from parsed markdown table rows."""
    if not rows or not rows[0]:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=0, cols=ncols)
    table.style = 'Table Grid'

    # Header row
    hdr = table.add_row()
    for ci, cell_text in enumerate(rows[0]):
        cell = hdr.cells[ci]
        cell.text = ''
        _shade_cell(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        _run(p, cell_text.strip(), size=SMALL_SZ, bold=True, color=WHITE)

    # Data rows
    for ri, row_data in enumerate(rows[1:]):
        row = table.add_row()
        for ci, cell_text in enumerate(row_data):
            if ci >= ncols:
                break
            cell = row.cells[ci]
            cell.text = ''
            if ri % 2 == 1:
                _shade_cell(cell, TABLE_ALT_BG)
            p = cell.paragraphs[0]
            text = cell_text.strip()
            # Color status cells
            color = BLACK
            if text in ('PASS', 'GREEN', 'green'):
                color = GREEN
            elif text in ('RED', 'FAIL', 'red'):
                color = RED
            elif text in ('UNKNOWN', 'DEGRADED', 'yellow', 'YELLOW'):
                color = ORANGE
            _run(p, text, size=SMALL_SZ, color=color)


def _parse_table_block(lines, start):
    """Parse a markdown table starting at index `start`. Returns (rows, end_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i


def _status_color(text):
    if any(w in text for w in ('RED', 'FAIL', 'FAILED')):
        return RED
    if any(w in text for w in ('DEGRADED', 'UNKNOWN', 'YELLOW', 'PARTIAL')):
        return ORANGE
    if any(w in text for w in ('GREEN', 'PASS', 'all green')):
        return GREEN
    return BLACK


def convert(md_path: Path, out_path: Path) -> Path:
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = BODY_SZ

    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    # -- Extract date from content ------------------------------------------
    date_str = datetime.now().strftime('%A, %B %d, %Y')
    for line in lines[:6]:
        m = re.search(r'(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s+(\w+ \d+,?\s*\d{4})', line)
        if m:
            date_str = f"{m.group(1)}, {m.group(2)}"
            break
        m = re.search(r'(\d{4}-\d{2}-\d{2})', line)
        if m:
            try:
                d = datetime.strptime(m.group(1), '%Y-%m-%d')
                date_str = d.strftime('%A, %B %d, %Y')
                break
            except ValueError:
                pass

    # -- Title block --------------------------------------------------------
    _para(doc, 'Daily Executive Briefing', TITLE_SZ, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    _para(doc, f'{date_str}  |  Prepared by Amy',
          SUBTITLE_SZ, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))

    llm_line = ''
    for line in lines[:10]:
        if line.strip().startswith('LLM'):
            llm_line = line.strip()
            break
    if llm_line:
        _para(doc, llm_line, SUBTITLE_SZ, italic=True, color=GRAY,
              align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(12))

    # -- Skip raw header up to first --- ------------------------------------
    i = 0
    while i < len(lines):
        if lines[i].strip() == '---':
            i += 1
            break
        i += 1

    # -- Process body -------------------------------------------------------
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank
        if not stripped:
            i += 1
            continue

        # HR
        if stripped == '---':
            i += 1
            continue

        # Markdown heading ## SECTION N: ...
        m = re.match(r'^#{1,3}\s+(.*)', stripped)
        if m:
            heading = m.group(1)
            _para(doc, heading, SECTION_SZ, bold=True, color=NAVY,
                  before=Pt(16), after=Pt(6))
            i += 1
            continue

        # Markdown table
        if stripped.startswith('|') and '|' in stripped[1:]:
            rows, i = _parse_table_block(lines, i)
            if rows:
                _add_table(doc, rows)
            continue

        # ALL-CAPS section headers (flat format from manual-briefing-v3.js)
        caps_headers = [
            'ACTION ITEMS', 'OPEN COMMITMENTS', 'INNER CIRCLE',
            'AI & TECH NEWS', 'WORLD NEWS', 'INDUSTRY NEWS',
            'FEATURE BACKLOG', 'OVERNIGHT ENHANCEMENTS',
            'LINKEDIN NETWORK INTELLIGENCE', 'COMMUNICATIONS SUMMARY',
            'CONTENT PIPELINE', 'SYSTEM HEALTH', 'REPUTATION MENTIONS',
            'WEEKLY THEME BRIEFING', 'PROJECTC INVOICE ACTIVITY',
        ]
        is_section = any(stripped.startswith(h) for h in caps_headers)
        if is_section:
            _para(doc, stripped, SECTION_SZ, bold=True, color=NAVY,
                  before=Pt(16), after=Pt(6))
            i += 1
            continue

        # Numbered article title
        if re.match(r'^\d+\.\s+[A-Z]', stripped) and len(stripped) > 20:
            _para(doc, stripped, ARTICLE_SZ, bold=True, color=BLUE,
                  before=Pt(10), after=Pt(2))
            i += 1
            continue

        # Source / URL lines
        if stripped.startswith('http') or stripped.startswith('Source:'):
            _para(doc, stripped, SOURCE_SZ, italic=True, color=GRAY, after=Pt(4))
            i += 1
            continue
        if re.match(r'^\s*(TechCrunch|BBC|NYT|NPR|AP|Reuters|Ars Technica|The Verge|Google News)', stripped):
            _para(doc, stripped, SOURCE_SZ, italic=True, color=GRAY, after=Pt(4))
            i += 1
            continue
        if re.match(r'^\s*https?://', stripped):
            _para(doc, stripped, SOURCE_SZ, italic=True, color=GRAY, after=Pt(4))
            i += 1
            continue

        # Bold markdown (**text**)
        if stripped.startswith('**') and stripped.endswith('**'):
            inner = stripped.strip('*').strip()
            _para(doc, inner, BODY_SZ, bold=True, color=BLACK, before=Pt(4), after=Pt(2))
            i += 1
            continue

        # Status lines with checkmarks
        if stripped.startswith(('✓', '✗', '?', '!')):
            color = GREEN if stripped.startswith('✓') else RED if stripped.startswith('✗') else BLACK
            _para(doc, stripped, BODY_SZ, color=color, before=Pt(1), after=Pt(1))
            i += 1
            continue

        # Overall health / status summary lines
        if 'Overall' in stripped and ('Health' in stripped or 'RED' in stripped or 'GREEN' in stripped or 'green' in stripped):
            color = _status_color(stripped)
            _para(doc, stripped, BODY_SZ, bold=True, color=color, before=Pt(4), after=Pt(4))
            i += 1
            continue

        # Bullet points
        if re.match(r'^\s*[*\-]\s', stripped) or re.match(r'^\s+[*\-]\s', line):
            clean = re.sub(r'^[\s*\-]+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            _run(p, clean, size=BODY_SZ)
            i += 1
            continue

        # Indented detail (2+ leading spaces)
        if line.startswith('  ') and not line.startswith('    '):
            color = BLACK
            if stripped.startswith('✓'):
                color = GREEN
            elif stripped.startswith(('✗', '!')):
                color = RED
            _para(doc, stripped, BODY_SZ, color=color, before=Pt(1), after=Pt(1))
            i += 1
            continue

        # Deep indent
        if line.startswith('    ') or line.startswith('\t'):
            _para(doc, stripped, SMALL_SZ, color=GRAY, before=Pt(0), after=Pt(1))
            i += 1
            continue

        # Footer
        if 'Reply with questions' in stripped:
            _para(doc, stripped, BODY_SZ, italic=True, color=GRAY, before=Pt(12))
            i += 1
            continue

        # Default body
        _para(doc, stripped, BODY_SZ, color=BLACK, before=Pt(2), after=Pt(2))
        i += 1

    # -- Save ---------------------------------------------------------------
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        fallback = out_path.with_name(out_path.stem + '.new' + out_path.suffix)
        doc.save(fallback)
        print(f'warning: {out_path} locked, wrote {fallback} instead', file=sys.stderr)
        return fallback


def main():
    if len(sys.argv) != 3:
        print('Usage: briefing-to-docx.py <input.md> <output.docx>', file=sys.stderr)
        sys.exit(2)
    actual = convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f'wrote {actual}')


if __name__ == '__main__':
    main()
